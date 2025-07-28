from docx import Document
from datetime import datetime
import os, sys

# Gets the correct path to the template file when running as an .exe or as a python file
def resource_path(relative_path):
    # Get the path to the resource (file) bundled with the executable
    try:
        # PyInstaller creates a temp folder and stores the path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

# Save the modified document to a new file
def save_document(doc, name):
    doc.save(f'{name}.docx')
    return f"{name}.docx"

# Replaces words in the body of the document
def replace_document_words(hashmap, doc):            
    placeholder_mapping = {
        'DEVICE_PH': 'deviceType',
        'CUSTOMER_PH': 'customer',
        'SERIAL_PH': 'serial',
        'MID_PH': 'mid',
        'MODEM_PH': 'modem',
        'IMEI_PH': 'imei',
        'PI_PH': 'pi_version',
        'TESTFIXTURE_PH': 'testFixer',
        'VOLTAGETESTER_PH': 'voltageTester',
        'BP_PH': 'batteryPack',
        'USB_PH': 'usb',
        'CAM_PH': 'cameraType',
        'CAMSN_PH': 'cameraSerial',
        'USER_PH': 'cameraUsername',
        'PASS_PH': 'cameraPass',
        'IP_PH': 'cameraIP',
    }
    
    # Iterate through each paragraph in the document
    for paragraph in doc.paragraphs:
        # Iterate through each run (a part of the paragraph with the same style/formatting)
        for run in paragraph.runs:
            for placeholder, key in placeholder_mapping.items():
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, hashmap[key])
                    
# replaces words in tables
def replace_table_words(hashmap, doc):
    table_mapping = {
        'MODEM_PH': 'modem',
        'SD_PH': 'sd_size',
        'IMEI_PH': 'imei',
        'PI_PH': 'pi_version',
        'COMPRESSION_TYPE_PH': 'compression_type',
        'TESTFIXTURE_PH': 'testFixer',
        'VOLTAGETESTER_PH': 'voltageTester',
        'BP_PH': 'batteryPack',
        'USB_PH': 'usb',
        'CAM_PH': 'cameraType',
        'CAMSN_PH': 'cameraSerial',
        'USER_PH': 'cameraUsername',
        'PASS_PH': 'cameraPass',
        'IP_PH': 'cameraIP',
    }
    
    # Iterate through tables in the document
    for table in doc.tables:
        # Iterate through each row in table
        for row in table.rows:
            # Iterate through each column in the row
            for cell in row.cells:
                # Iterate through each paragraph in the cell (cells contain paragraphs)
                for paragraph in cell.paragraphs:
                    # Iterate through each word in the paragraph
                    for run in paragraph.runs:
                        for placeholder, key in table_mapping.items():
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, hashmap[key])

# In the created document this section replaces words in the header
def replace_header_words(hashmap, doc):
    for section in doc.sections:
        header = section.header

        for paragraph in header.paragraphs:
            for run in paragraph.runs:
                if 'DATE_PH' in run.text:
                    current_date = datetime.now().strftime("%m/%d/%Y")
                    run.text = run.text.replace('DATE_PH', current_date)
                    continue

def generate_checklist_document(hashmap):
    documentLocation = None
    deviceType = hashmap["deviceType"]

    # Load the template document
    if deviceType == "LHC/LHG":
        doc = Document(resource_path('templates/LHCTemplate.docx'))  # Use resource_path to get the bundled file
    else:
        doc = Document(resource_path(f'templates/{deviceType}Template.docx'))
        
    replace_document_words(hashmap, doc)
    replace_table_words(hashmap, doc)
    
    # Save the document and return the location
    documentLocation = save_document(doc, hashmap["serial"])

    return documentLocation

def generate_customer_info_document(hashmap):
    documentLocation = None
    deviceType = hashmap["deviceType"]

    # Load the template document
    if deviceType == "LHC/LHG":
        doc = Document(resource_path('customerInfoTemplates/LHCCustomerTemplate.docx'))  # Use resource_path to get the bundled file
    else:
        doc = Document(resource_path(f'customerInfoTemplates/{deviceType}CustomerTemplate.docx'))
        
    replace_header_words(hashmap, doc)
    
    replace_document_words(hashmap, doc)
    
    documentLocation = save_document(doc, f'{hashmap["serial"]}_customer')

    return documentLocation
